# -*- coding: utf-8 -*-
"""better_docx.helpers 测试。"""
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from better_docx import (
    set_font, add_para, shade_cell, set_cell_text, add_para_to_cell,
    new_image_paragraph, add_image, global_replace, set_cell_widths,
    set_repeat_header, set_row_no_split, diagnose_pagination,
    fix_one_paragraph_per_page,
)


def test_set_font_sets_eastasia():
    d = Document()
    r = d.add_paragraph().add_run("中文")
    set_font(r, size=12, bold=True, italic=True, color=(0xFF, 0, 0))
    assert r._element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
    assert r.italic and r.bold


def test_add_para_italic():
    d = Document()
    add_para(d, "题注", size=10, italic=True, color=(0x60, 0x60, 0x60))
    assert d.paragraphs[0].runs[0].italic is True


def test_set_cell_text_single_run():
    d = Document()
    c = d.add_table(1, 1).rows[0].cells[0]
    set_cell_text(c, "内容", bold=True)
    assert c.text == "内容"
    assert c.paragraphs[0].runs[0].bold
    assert len(c.paragraphs[0].runs) == 1          # 不应残留空 run


def test_add_para_to_cell_appends():
    d = Document()
    c = d.add_table(1, 1).rows[0].cells[0]
    set_cell_text(c, "标题")
    add_para_to_cell(c, "正文")
    assert "标题" in c.text and "正文" in c.text
    assert len(c.paragraphs) == 2


def test_new_image_paragraph_attached_to_body():
    d = Document()
    p = new_image_paragraph(d)
    assert p._p.getparent() is d.element.body


def test_add_image_missing_returns_none():
    d = Document()
    assert add_image(d, r"D:\不存在.png", caption="x") is None


def test_global_replace_covers_table():
    d = Document()
    d.add_paragraph("价格 100 元")
    t = d.add_table(1, 1)
    t.rows[0].cells[0].text = "价格 100 元"
    n = global_replace(d, "100", "200")
    assert n == 2
    assert "200" in d.paragraphs[0].text
    assert "200" in t.rows[0].cells[0].text


def test_set_cell_widths():
    d = Document()
    t = d.add_table(1, 3)
    set_cell_widths(t, [2.0, 3.0, 4.0])
    assert t.autofit is False
    assert abs(t.rows[0].cells[2].width.cm - 4.0) < 0.01


def test_repeat_header_and_no_split():
    d = Document()
    t = d.add_table(2, 2)
    set_repeat_header(t)
    set_row_no_split(t.rows[0])
    trPr = t.rows[0]._tr.trPr
    assert trPr.find(qn("w:tblHeader")) is not None
    assert trPr.find(qn("w:cantSplit")) is not None


def test_diagnose_and_fix_pagination():
    d = Document()
    p1 = d.add_paragraph()
    r1 = p1.add_run("关键文字不丢失")
    br1 = OxmlElement("w:br")
    br1.set(qn("w:type"), "page")
    r1._element.append(br1)
    p2 = d.add_paragraph("段前分页段")
    p2.paragraph_format.page_break_before = True
    cell = d.add_table(1, 1).rows[0].cells[0]
    cell.paragraphs[0].add_run("表格文字")
    br3 = OxmlElement("w:br")
    br3.set(qn("w:type"), "page")
    cell.paragraphs[0].add_run()._element.append(br3)

    before = diagnose_pagination(d)
    assert before["page_break_before_paragraphs"] == 1
    assert before["hard_page_breaks"] == 2
    assert before["hard_page_breaks_in_tables"] == 1

    stats = fix_one_paragraph_per_page(d)
    assert stats == {"page_break_before": 1, "hard_page_breaks": 2}

    after = diagnose_pagination(d)
    assert after["page_break_before_paragraphs"] == 0
    assert after["hard_page_breaks"] == 0
    assert "关键文字不丢失" in p1.text          # 文字不被 run.clear() 误删
    assert "表格文字" in cell.paragraphs[0].text