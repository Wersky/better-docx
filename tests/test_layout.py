# -*- coding: utf-8 -*-
"""better_docx.layout 测试。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from better_docx import (
    set_page_a4, set_orientation, enable_even_odd_headers,
    add_page_number_field, add_numpages_field, set_first_line_indent_chars,
    set_line_spacing, add_bullets, add_numbered, restyle_heading, add_toc_field,
)


def test_set_page_a4():
    d = Document()
    s = d.sections[0]
    set_page_a4(s, (2.0, 2.0, 3.0, 3.0))
    assert abs(s.page_width.cm - 21.0) < 0.01
    assert abs(s.top_margin.cm - 2.0) < 0.01


def test_set_orientation_idempotent():
    d = Document()
    s = d.sections[0]
    set_page_a4(s)
    set_orientation(s, True)
    set_orientation(s, True)          # 幂等
    assert s.page_width.cm > s.page_height.cm
    set_orientation(s, False)
    assert s.page_width.cm < s.page_height.cm


def test_enable_even_odd_headers():
    d = Document()
    s = d.sections[0]
    enable_even_odd_headers(s)
    enable_even_odd_headers(s)        # 幂等，不重复
    names = [c.tag.split("}")[-1] for c in s._sectPr]
    assert names.count("evenAndOddHeaders") == 1


def test_page_number_field():
    d = Document()
    s = d.sections[0]
    add_page_number_field(s.footer.paragraphs[0], before="第 ", after=" 页")
    add_numpages_field(s.footer.paragraphs[0])
    instrs = [t.text.strip() for t in s.footer.paragraphs[0]._element.findall(".//" + qn("w:instrText"))]
    flds = [f.get(qn("w:fldCharType")) for f in s.footer.paragraphs[0]._element.findall(".//" + qn("w:fldChar"))]
    assert instrs == ["PAGE", "NUMPAGES"]
    assert flds == ["begin", "separate", "end", "begin", "separate", "end"]


def test_page_number_roman():
    d = Document()
    s = d.sections[0]
    add_page_number_field(s.footer.paragraphs[0], num_fmt="ROMAN")
    instr = s.footer.paragraphs[0]._element.findall(".//" + qn("w:instrText"))[0].text
    assert "ROMAN" in instr and "PAGE" in instr


def test_first_line_indent_chars():
    d = Document()
    p = d.add_paragraph("x")
    set_first_line_indent_chars(p, chars=2)
    ind = p._p.pPr.find(qn("w:ind"))
    assert ind.get(qn("w:firstLineChars")) == "200"
    assert ind.get(qn("w:firstLine")) == "480"


def test_first_line_indent_half_char():
    d = Document()
    p = d.add_paragraph("x")
    set_first_line_indent_chars(p, chars=0.5)
    assert p._p.pPr.find(qn("w:ind")).get(qn("w:firstLineChars")) == "50"


def test_line_spacing():
    d = Document()
    p1 = d.add_paragraph("a")
    set_line_spacing(p1, 1.5)
    p2 = d.add_paragraph("b")
    set_line_spacing(p2, Pt(20), rule=WD_LINE_SPACING.EXACTLY)
    assert p1.paragraph_format.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE
    assert p2.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY


def test_lists():
    d = Document()
    add_bullets(d, ["甲", "乙"])
    add_numbered(d, [("步骤1", ["子1", "子2"]), ("步骤2", None)])
    styles = [p.style.name for p in d.paragraphs]
    assert "List Bullet" in styles
    assert "List Number 2" in styles


def test_restyle_heading():
    d = Document()
    restyle_heading(d.styles["Heading 1"], "黑体", 16, (0, 0, 0))
    h = d.styles["Heading 1"]
    assert h.font.name == "黑体"
    assert h.font.color.rgb == RGBColor(0, 0, 0)


def test_add_toc_field():
    d = Document()
    d.add_paragraph("前置")
    add_toc_field(d)
    instrs = [t.text for t in d.element.body.findall(".//" + qn("w:instrText"))]
    assert "TOC" in "".join(instrs)
    assert len(d.element.body.findall(".//" + qn("w:sdt"))) == 1