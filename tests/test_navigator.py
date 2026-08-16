# -*- coding: utf-8 -*-
"""better_docx.navigator 测试。"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from better_docx import (
    docx_to_json, docx_to_markdown, build_heading_index, find_heading,
    find_paragraph_containing, insert_paragraph_after, insert_paragraph_before_cjk,
    insert_table_before, insert_table_after, list_bookmarks,
    find_bookmark_paragraph, append_text_to_bookmark, add_bookmark,
    add_hyperlink, list_hyperlinks,
)


def test_docx_to_json_outline_level():
    d = Document()
    d.add_paragraph("普通段落")
    p = d.add_paragraph("带大纲级别段落")
    pPr = p._p.get_or_add_pPr()
    oulvl = OxmlElement("w:outlineLvl")
    oulvl.set(qn("w:val"), "2")
    pPr.append(oulvl)
    assert docx_to_json(d)["blocks"][1]["outline_level"] == 2


def test_docx_to_json_block_types():
    d = Document()
    d.add_heading("标题", 1)
    d.add_paragraph("正文")
    t = d.add_table(1, 2)
    t.rows[0].cells[0].text = "a"
    types = [b["type"] for b in docx_to_json(d)["blocks"]]
    assert types == ["paragraph", "paragraph", "table", "sectPr"]


def test_docx_to_markdown():
    d = Document()
    d.add_heading("第一章", 1)
    d.add_heading("1.1 子节", 2)
    d.add_paragraph("正文")
    md = docx_to_markdown(d)
    assert "# 第一章" in md and "## 1.1 子节" in md


def test_docx_to_markdown_multilink():
    d = Document()
    p = d.add_paragraph()
    add_hyperlink(p, "https://a.com", "链接A")
    p.add_run(" 和 ")
    add_hyperlink(p, "https://b.com", "链接B")
    md = docx_to_markdown(d)
    assert "https://a.com" in md and "https://b.com" in md


def test_build_heading_index_with_custom_style():
    d = Document()
    st = d.styles.add_style("Heading 自定义", WD_STYLE_TYPE.PARAGRAPH)
    d.add_heading("正常标题", 1)
    p = d.add_paragraph("自定义标题段")
    p.style = d.styles["Heading 自定义"]
    idx = build_heading_index(d)      # 不应因 "自定义" 非数字而崩溃
    assert "正常标题" in idx


def test_find_heading():
    d = Document()
    d.add_heading("第三章 结论", 1)
    d.add_heading("3.1 子节", 2)
    assert find_heading(d, "第三章 结论") is not None
    assert find_heading(d, "结论", exact=False).text == "第三章 结论"


def test_find_paragraph_containing():
    d = Document()
    d.add_paragraph("数据预处理流程")
    assert find_paragraph_containing(d, "预处理") is not None


def test_insert_paragraph_after():
    d = Document()
    p = d.add_paragraph("目标")
    np = insert_paragraph_after(p, "新增")
    assert p._p.getnext() is np._p


def test_insert_paragraph_before_cjk():
    d = Document()
    p = d.add_paragraph("目标")
    np = insert_paragraph_before_cjk(p, "中文前段")
    assert np._p.getnext() is p._p
    assert np.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"


def test_insert_table_before():
    d = Document()
    p = d.add_paragraph("锚")
    tbl = insert_table_before(p, rows=2, cols=3, caption="表",
                              fill=lambda r, c: f"{r}{c}")
    assert tbl.rows[0].cells[0].text == "00"


def test_insert_table_after():
    d = Document()
    p = d.add_paragraph("锚")
    tbl = insert_table_after(p, rows=2, cols=2)
    assert p._p.getnext() is tbl._tbl


def test_bookmarks():
    d = Document()
    p = d.add_paragraph("正文内容")
    nid = add_bookmark(p, "标记", 3)
    assert nid == 4
    assert list_bookmarks(d) == [{"name": "标记", "id": "3"}]
    para, bm = find_bookmark_paragraph(d, "标记")
    assert para is not None
    assert append_text_to_bookmark(d, "标记", "-追加")
    assert "-追加" in para.text


def test_hyperlink():
    d = Document()
    add_hyperlink(d.add_paragraph(), "https://x.com", "链接")
    assert list_hyperlinks(d) == [("链接", "https://x.com")]