# -*- coding: utf-8 -*-
"""示例 1：生成带原生 Word 表格的成绩单（合并单元格、底纹、列宽控制）。
运行：python examples/make_table.py
输出：examples/表格示例.docx
"""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from better_docx import set_font, shade_cell, set_cell_text, set_cell_widths, add_para

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "表格示例.docx")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("2026 春季学期《机器学习》成绩单"), name="黑体", size=16, bold=True)
add_para(doc, "数据系 · 计算机与软件学院   |   制表日期：2026-08-16",
         size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x80, 0x80, 0x80), line_spacing=1.0)

# 原生表格（Table Grid 边框样式）
table = doc.add_table(rows=8, cols=5)
table.style = "Table Grid"
set_cell_widths(table, [2.6, 3.2, 3.2, 3.2, 3.3])   # 合计 15.5cm ≤ 正文宽

# 第 0 行：整行合并标题 + 深蓝底白字
merged = table.rows[0].cells[0].merge(table.rows[0].cells[-1])
set_cell_text(merged, "成绩汇总表", bold=True, size=12, color=RGBColor(0xFF, 0xFF, 0xFF))
shade_cell(merged, "4472C4")

# 第 1 行：表头 + 浅蓝底
for i, h in enumerate(["姓名", "学号", "平时(30%)", "期末(70%)", "总评"]):
    set_cell_text(table.rows[1].cells[i], h, bold=True)
    shade_cell(table.rows[1].cells[i], "D9E2F3")

# 第 2~6 行：数据
data = [
    ["张伟", "2024156001", "88", "92", "优秀"],
    ["李娜", "2024156002", "75", "81", "良好"],
    ["王强", "2024156003", "63", "70", "及格"],
    ["赵敏", "2024156004", "95", "98", "优秀"],
    ["刘政", "2024156012", "91", "89", "优秀"],
]
for ri, row_data in enumerate(data, start=2):
    for ci, val in enumerate(row_data):
        set_cell_text(table.rows[ri].cells[ci], val)

# 第 7 行：前两列合并的"平均分"行 + 浅灰底
merged = table.rows[7].cells[0].merge(table.rows[7].cells[1])
set_cell_text(merged, "平均分", bold=True)
for ci, v in [(2, "82.4"), (3, "86.0"), (4, "—")]:
    set_cell_text(table.rows[7].cells[ci], v, bold=True)
for c in table.rows[7].cells:
    shade_cell(c, "F2F2F2")

doc.save(OUT)
print("已生成:", OUT)