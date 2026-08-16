# -*- coding: utf-8 -*-
"""示例 2：多节研究报告（封面 → 目录 → 正文 → 横向附录 → 竖向附录）。

覆盖：分节横竖混排、页眉页脚/首页不同/奇偶不同、PAGE/NUMPAGES 域、
TOC 目录域、中文缩进、列表、表格、图片、书签、超链接、导航定位与定点注入。

运行：python examples/make_report.py
输出：examples/研究报告.docx
"""
import os
import sys
import struct
import zlib
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from better_docx import (
    set_font, set_cell_text, set_cell_widths, shade_cell, add_image,
    set_page_a4, set_orientation, enable_even_odd_headers,
    add_page_number_field, add_numpages_field, set_first_line_indent_chars,
    restyle_heading, add_toc_field, add_bookmark, add_hyperlink,
    find_heading, insert_table_before, docx_to_markdown,
)

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "研究报告.docx")
CHART = os.path.join(os.path.dirname(os.path.abspath(__file__)), "growth.png")


def make_png(path, w=420, h=140, color=(0x2F, 0x54, 0x96)):
    """生成一张纯色 PNG（无 Pillow 依赖），用于示例插图。"""
    def chunk(tag, data):
        c = struct.pack(">I", len(data)) + tag + data
        return c + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)

    raw = (b"\x00" + bytes(color) * w) * h
    png = (b"\x89PNG\r\n\x1a\n"
           + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
           + chunk(b"IDAT", zlib.compress(raw))
           + chunk(b"IEND", b""))
    with open(path, "wb") as f:
        f.write(png)


make_png(CHART)

doc = Document()
# 全局样式：正文宋体 + 标题黑体
normal = doc.styles["Normal"]
normal.font.name = "宋体"
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.paragraph_format.line_spacing = 1.5
restyle_heading(doc.styles["Heading 1"], "黑体", 16, (0, 0, 0))

add_toc_field(doc)                      # 目录域（最前，Word 打开 F9 更新）

# ---- 节 1：封面 + 正文（首页不同、奇偶不同） ----
sec1 = doc.sections[0]
set_page_a4(sec1)
set_orientation(sec1, landscape=False)
sec1.different_first_page_header_footer = True
enable_even_odd_headers(sec1)
sec1.header.is_linked_to_previous = False
sec1.header.paragraphs[0].text = "2026 市场研究报告"
sec1.even_page_header.is_linked_to_previous = False
sec1.even_page_header.paragraphs[0].text = "偶数页 · 2026 市场研究报告"
sec1.footer.is_linked_to_previous = False
add_page_number_field(sec1.footer.paragraphs[0], before="第 ", after=" 页")
add_numpages_field(sec1.footer.paragraphs[0])

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("2026 年市场研究报告"), name="黑体", size=26, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("数据驱动 · 竞争分析 · 趋势研判"), size=14, color=(0x60, 0x60, 0x60))
doc.add_page_break()

doc.add_heading("一、市场概览", level=1)
p = doc.add_paragraph("2026 年上半年市场规模同比增长 12.4%，增速高于去年同期。本文基于公开数据构建分析框架。")
set_first_line_indent_chars(p, chars=2)
for it in ["市场规模 428 亿元", "用户渗透率 67%", "集中度 CR5 = 38%"]:
    doc.add_paragraph(it, style="List Bullet")

doc.add_heading("二、竞争分析", level=1)
t = doc.add_table(rows=3, cols=4)
t.style = "Table Grid"
set_cell_widths(t, [3.0, 3.0, 3.0, 3.0])
for i, h in enumerate(["厂商", "份额", "增速", "备注"]):
    set_cell_text(t.rows[0].cells[i], h, bold=True)
    shade_cell(t.rows[0].cells[i], "D9E2F3")
for ri, row in enumerate([["A 公司", "15%", "+9%", "龙头"],
                          ["B 公司", "12%", "+21%", "增长最快"]], start=1):
    for ci, v in enumerate(row):
        set_cell_text(t.rows[ri].cells[ci], v)

add_image(doc, CHART, width_inches=5.2, caption="图 2-1 市场规模增长趋势")

p_end = doc.add_paragraph("详细数据见附录。")
add_bookmark(p_end, "附录指引", bookmark_id=7)
p_link = doc.add_paragraph("更多信息请访问：")
add_hyperlink(p_link, "https://example.com/report2026", "官方报告链接")

doc.add_heading("三、结论", level=1)
doc.add_paragraph("结论：市场整体向上，建议关注 B 公司所在细分赛道。")

# ---- 节 2：横向附录（宽表格） ----
sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
set_page_a4(sec2)
set_orientation(sec2, landscape=True)
sec2.header.is_linked_to_previous = False
sec2.header.paragraphs[0].text = "附录 A · 横向数据表"
sec2.footer.is_linked_to_previous = False
add_page_number_field(sec2.footer.paragraphs[0], before="附A 第 ", after=" 页")
doc.add_heading("附录 A 厂商全量数据", level=1)
t2 = doc.add_table(rows=3, cols=8)
t2.style = "Table Grid"
set_cell_widths(t2, [2.0] * 8)
for ci, h in enumerate(["厂商", "营收(亿)", "利润(亿)", "份额", "增速", "研发占比", "海外占比", "评级"]):
    set_cell_text(t2.rows[0].cells[ci], h, bold=True)
    shade_cell(t2.rows[0].cells[ci], "D9E2F3")
for ri in (1, 2):
    for ci in range(8):
        set_cell_text(t2.rows[ri].cells[ci], f"{ri}{ci}9.9")

# ---- 节 3：回竖版 ----
sec3 = doc.add_section(WD_SECTION.NEW_PAGE)
set_page_a4(sec3)
set_orientation(sec3, landscape=False)
sec3.header.is_linked_to_previous = False
sec3.header.paragraphs[0].text = "附录 B"
sec3.footer.is_linked_to_previous = False
add_page_number_field(sec3.footer.paragraphs[0], before="附B 第 ", after=" 页")
doc.add_heading("附录 B 方法论说明", level=1)
doc.add_paragraph("数据来源与估算方法的说明文字。")

doc.save(OUT)
print("已生成:", OUT, "| 节数:", len(doc.sections))

# ---- 展示 navigator：定位 + 定点注入 + 解析回读 ----
d2 = Document(OUT)
assert find_heading(d2, "三、结论") is not None
insert_table_before(find_heading(d2, "三、结论"), rows=2, cols=3,
                    caption="注入表：风险矩阵", fill=lambda r, c: f"R{r}C{c}")
d2.save(OUT.replace(".docx", "_注入演示.docx"))
print("注入演示已生成:", OUT.replace(".docx", "_注入演示.docx"))
print("--- Markdown 预览（前 8 行）---")
print("\n".join(docx_to_markdown(d2).splitlines()[:8]))