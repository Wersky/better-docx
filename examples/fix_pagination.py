# -*- coding: utf-8 -*-
"""示例 3：分页异常（一页一段）的诊断与修复。

演示如何识别并修复「长文本插入后一段独占一页、文档张数暴涨」的问题。

运行：python examples/fix_pagination.py
"""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from better_docx import diagnose_pagination, fix_one_paragraph_per_page

DIR = os.path.dirname(os.path.abspath(__file__))
DIRTY = os.path.join(DIR, "病态文档.docx")
CLEAN = os.path.join(DIR, "修复后文档.docx")


def make_dirty(path):
    """构造一个带分页异常的文档：段前分页 + 文字/分页符同 run + 表格内分页。"""
    doc = Document()
    p1 = doc.add_paragraph()
    r1 = p1.add_run("这段文字绝不能丢")
    br1 = OxmlElement("w:br")
    br1.set(qn("w:type"), "page")          # 硬分页符和文字挤在同一 run
    r1._element.append(br1)

    p2 = doc.add_paragraph("段前分页段落")
    p2.paragraph_format.page_break_before = True   # 误设段前分页

    t = doc.add_table(1, 1)                 # 表格内的分页符（doc.paragraphs 遍历不到）
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].add_run("表格内文字")
    br3 = OxmlElement("w:br")
    br3.set(qn("w:type"), "page")
    cell.paragraphs[0].add_run()._element.append(br3)

    doc.save(path)


make_dirty(DIRTY)
doc = Document(DIRTY)

print("修复前诊断:", diagnose_pagination(doc))
stats = fix_one_paragraph_per_page(doc)
print("修复统计:", stats)
doc.save(CLEAN)

d2 = Document(CLEAN)
print("修复后诊断:", diagnose_pagination(d2))
assert diagnose_pagination(d2)["hard_page_breaks"] == 0
assert "这段文字绝不能丢" in d2.paragraphs[0].text, "文字被误删！"
print("关键文字保留 ✅  已生成:", CLEAN)