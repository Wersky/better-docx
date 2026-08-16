# -*- coding: utf-8 -*-
"""
better_docx.layout —— 页面布局与排版：分节/页眉页脚/页码域、段落缩进/列表、
全局样式与自动目录（TOC 域）。
"""
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ================= 子功能 6：全局布局与结构 =================

# CT_SectPr 子元素 schema 顺序（ECMA-376）
_SECTPR_ORDER = ["footnotePr", "endnotePr", "type", "pgSz", "pgMar", "paperSrc",
                 "pgBorders", "lnNumType", "pgNumType", "cols", "formProt",
                 "vAlign", "noEndnote", "titlePg", "evenAndOddHeaders",
                 "textDirection", "bidi", "rtlGutter", "docGrid",
                 "printerSettings", "sectPrChange"]


def _insert_schema_ordered(sectPr, el):
    """按 sectPr schema 顺序插入子元素（Word 宽容但 LibreOffice/校验器可能严格）。"""
    name = el.tag.split("}")[-1]
    idx = _SECTPR_ORDER.index(name)
    for child in reversed(list(sectPr)):
        cname = child.tag.split("}")[-1]
        if cname in _SECTPR_ORDER and _SECTPR_ORDER.index(cname) < idx:
            child.addnext(el)
            return
    sectPr.insert(0, el)


def set_page_a4(section, margin_cm=(2.54, 2.54, 3.17, 3.17)):
    """A4 纸张 + 页边距（上/下/左/右，cm）。默认 Word 中文常用边距。"""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(margin_cm[0])
    section.bottom_margin = Cm(margin_cm[1])
    section.left_margin = Cm(margin_cm[2])
    section.right_margin = Cm(margin_cm[3])


def set_orientation(section, landscape=False):
    """设置节方向，自动交换宽高。幂等。
    add_section() 会复制上一节 sectPr（含 orient/titlePg 等属性），只改宽高不设
    orientation 会造成"横向标志+竖版尺寸"的不一致状态。"""
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    w, h = section.page_width, section.page_height
    if landscape and w < h:
        section.page_width, section.page_height = h, w
    elif not landscape and w > h:
        section.page_width, section.page_height = h, w


def enable_even_odd_headers(section):
    """启用奇偶页不同（w:evenAndOddHeaders）。python-docx 无公开 API。
    警告：直接赋值 section.even_page_header_footer 会被静默忽略（动态属性），不生效！"""
    sectPr = section._sectPr
    if sectPr.find(qn("w:evenAndOddHeaders")) is None:
        _insert_schema_ordered(sectPr, OxmlElement("w:evenAndOddHeaders"))


def add_page_number_field(paragraph, before="第 ", after=" 页", num_fmt=None,
                          align=None):
    """在页眉/页脚段落插入 PAGE 动态页码域（复杂域三段式）。
    num_fmt: None=阿拉伯数字, 'ROMAN'/'roman'=大小写罗马, 'CHINESENUM1'=中文数字。
    同段追加总页数：再调 add_numpages_field()。"""
    if align is not None:
        paragraph.alignment = align
    fmt_txt = f" \\* {num_fmt}" if num_fmt else ""
    if before:
        paragraph.add_run(before)
    r_b = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r_b._element.append(fld)
    r_i = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGE {fmt_txt} "
    r_i._element.append(instr)
    r_s = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r_s._element.append(fld)
    paragraph.add_run("1")                      # 无 Word 时的占位显示
    r_e = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r_e._element.append(fld)
    if after:
        paragraph.add_run(after)


def add_numpages_field(paragraph, before="/ 共 ", after=" 页"):
    """在同一段落追加 NUMPAGES 域（总页数）。通常在 PAGE 域之后调用。"""
    if before:
        paragraph.add_run(before)
    r_b = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r_b._element.append(fld)
    r_i = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " NUMPAGES "
    r_i._element.append(instr)
    r_s = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r_s._element.append(fld)
    paragraph.add_run("1")
    r_e = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r_e._element.append(fld)
    if after:
        paragraph.add_run(after)


# ================= 子功能 7：段落与列表 =================

def set_first_line_indent_chars(paragraph, chars=2, font_size_pt=12):
    """中文首行缩进 N 字符。
    w:firstLineChars=200 → 2.00 字符（Word 按字符单位，随字号缩放）
    + w:firstLine 绝对 twips 兜底（LibreOffice 兼容）。不依赖空格/英制。"""
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(font_size_pt * chars)
    ind = paragraph._p.pPr.find(qn("w:ind"))
    if ind is not None:
        ind.set(qn("w:firstLineChars"), str(int(chars * 100)))


def set_line_spacing(paragraph, value, rule=None):
    """行距：float(1.0/1.5)=倍数；Pt(20)=固定值(exact)；显式 rule 可传
    WD_LINE_SPACING.SINGLE/ONE_POINT_FIVE/EXACTLY/AT_LEAST/MULTIPLE。"""
    pf = paragraph.paragraph_format
    if rule:
        pf.line_spacing_rule = rule
    pf.line_spacing = value
    return pf


def add_bullets(doc, items, style="List Bullet"):
    """项目符号段落（style 可选 'List Bullet'/'List Bullet 2'/'List Bullet 3'）。"""
    for it in items:
        doc.add_paragraph(it, style=style)


def add_numbered(doc, pairs, style="List Number", sub_style="List Number 2"):
    """编号列表 + 可选二级子项。pairs = [(一级文本, [子项...]|None), ...]。"""
    for text, subs in pairs:
        doc.add_paragraph(text, style=style)
        if subs:
            for s in subs:
                doc.add_paragraph(s, style=sub_style)


# ================= 子功能 8：全局样式与目录 =================

def restyle_heading(style_obj, name="黑体", size_pt=12, color_hex=None):
    """把内置 Heading 样式改成 黑体/字号/颜色，保留大纲级别。
    color_hex 为 (r,g,b) 元组或 None（不变色）。"""
    style_obj.font.name = name
    style_obj.font.size = Pt(size_pt)
    style_obj.font.bold = True
    if color_hex:
        from docx.shared import RGBColor
        style_obj.font.color.rgb = RGBColor(*color_hex)
    style_obj.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_toc_field(doc, include_levels="1-3",
                  placeholder="目录为空——在 Word 中右键此区域选择“更新域”"):
    """插入自动目录域（w:sdt 包装 TOC field），Word 打开 F9/右键更新生成目录。
    前提：正文标题必须用 Heading N 样式，否则目录无内容。插到文档最前。"""
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    docPartObj = OxmlElement("w:docPartObj")
    gallery = OxmlElement("w:docPartGallery")
    gallery.set(qn("w:val"), "Table of Contents")
    unique = OxmlElement("w:docPartUnique")
    docPartObj.append(gallery)
    docPartObj.append(unique)
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)

    content = OxmlElement("w:sdtContent")
    # 目录标题（居中黑体 16pt，不依赖模板样式）
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    p.append(pPr)
    r1 = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "32")                      # 16pt
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "黑体")
    rFonts.set(qn("w:eastAsia"), "黑体")
    rPr.append(rFonts)
    r1.append(rPr)
    t = OxmlElement("w:t")
    t.text = "目录"
    r1.append(t)
    p.append(r1)
    content.append(p)

    # 域代码段
    p2 = OxmlElement("w:p")
    r_b = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    fld.set(qn("w:dirty"), "true")                 # 打开提示更新
    r_b.append(fld)
    p2.append(r_b)
    r_i = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{include_levels}" \\h \\z \\u '
    r_i.append(instr)
    p2.append(r_i)
    r_s = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r_s.append(fld)
    p2.append(r_s)
    r_ph = OxmlElement("w:r")
    t_ph = OxmlElement("w:t")
    t_ph.text = placeholder
    r_ph.append(t_ph)
    p2.append(r_ph)
    r_e = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r_e.append(fld)
    p2.append(r_e)
    content.append(p2)
    sdt.append(content)

    doc.element.body[0].addprevious(sdt)

__all__ = [
    "set_page_a4", "set_orientation", "enable_even_odd_headers",
    "add_page_number_field", "add_numpages_field", "set_first_line_indent_chars",
    "set_line_spacing", "add_bullets", "add_numbered", "restyle_heading",
    "add_toc_field",
]
