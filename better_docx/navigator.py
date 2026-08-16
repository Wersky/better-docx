# -*- coding: utf-8 -*-
"""
better_docx.navigator —— 文档导航与状态感知：docx→JSON/Markdown 解析、锚点定位、
定点内容注入、书签与超链接读写。
"""
import copy
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ================= 解析：docx → JSON / Markdown =================

def _heading_level(style_name):
    """从样式名提取标题级别；非 'Heading N' 格式（如自定义 'Heading 自定义'）
    返回 None，避免 int('自定义') 崩溃。"""
    if not style_name or not style_name.startswith("Heading"):
        return None
    m = re.search(r"(\d+)$", style_name)
    return int(m.group(1)) if m else None


def _para_numpr(p):
    numPr = p._p.pPr.find(qn("w:numPr")) if p._p.pPr is not None else None
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    numId = numPr.find(qn("w:numId"))
    return {"ilvl": ilvl.get(qn("w:val")) if ilvl is not None else "0",
            "numId": numId.get(qn("w:val")) if numId is not None else None}


def _para_meta(p):
    """段落元信息：书签名 + 超链接[(text,target)]。"""
    bms = [b.get(qn("w:name")) for b in p._p.findall(qn("w:bookmarkStart"))]
    hls = []
    for hl in p._p.findall(qn("w:hyperlink")):
        rid = hl.get(qn("r:id"))
        target = None
        if rid:
            try:
                target = p.part.rels[rid].target_ref
            except KeyError:
                pass
        hls.append({"text": "".join(t.text or "" for t in hl.iter(qn("w:t"))),
                    "target": target})
    return bms, hls


def docx_to_json(doc):
    """docx（Document 对象或路径）→ JSON 树 dict。
    blocks: paragraph(text/style/outline_level/images/numPr/bookmarks/hyperlinks)
            / table(rows/cols/cells) / sectPr"""
    if isinstance(doc, str):
        doc = Document(doc)
    blocks = []
    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = Paragraph(child, doc)
            style = p.style.name if p.style else None
            outline = _heading_level(style)
            oulvl = child.find(f"{qn('w:pPr')}/{qn('w:outlineLvl')}")
            if oulvl is not None:
                outline = int(oulvl.get(qn("w:val")))
            bms, hls = _para_meta(p)
            blocks.append({
                "type": "paragraph", "text": p.text, "style": style,
                "outline_level": outline,
                "images": len(child.findall(".//" + qn("w:drawing"))),
                "numPr": _para_numpr(p), "bookmarks": bms, "hyperlinks": hls,
            })
        elif tag == "tbl":
            tbl = Table(child, doc)
            blocks.append({"type": "table", "rows": len(tbl.rows),
                           "cols": len(tbl.columns),
                           "cells": [[c.text for c in row.cells] for row in tbl.rows]})
        elif tag == "sectPr":
            blocks.append({"type": "sectPr"})
    return {"sections": len(doc.sections), "blocks": blocks}


def docx_to_markdown(doc):
    """docx → Markdown 文本：Heading→# 层级、表格→格线、链接→[]()、书签→锚点。"""
    if isinstance(doc, str):
        doc = Document(doc)
    lines = []
    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = Paragraph(child, doc)
            style = p.style.name if p.style else ""
            bms, hls = _para_meta(p)
            lv = _heading_level(style)
            if lv:
                lines.append("#" * lv + " " + p.text)
            elif hls:
                parts = []
                li = 0
                for r in child:
                    if r.tag == qn("w:hyperlink"):
                        htext = "".join(t.text or "" for t in r.iter(qn("w:t")))
                        target = hls[li]["target"] if li < len(hls) else None
                        parts.append(f"[{htext}]({target})")
                        li += 1
                    elif r.tag == qn("w:r"):
                        parts.append("".join(t.text or "" for t in r.findall(qn("w:t"))))
                lines.append("".join(parts))
            elif bms:
                lines.append(f"<a name={bms[0]}></a>{p.text}")
            else:
                lines.append(p.text)
        elif tag == "tbl":
            tbl = Table(child, doc)
            header = [c.text for c in tbl.rows[0].cells]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("|" + "---|" * len(header))
            for row in tbl.rows[1:]:
                lines.append("| " + " | ".join(c.text for c in row.cells) + " |")
    return "\n".join(l for l in lines if l.strip())


# ================= 锚点定位（Agent 之眼） =================

def build_heading_index(doc):
    """标题索引：{标题文本: [{"level": n, "paragraph": Paragraph}, ...]}"""
    if isinstance(doc, str):
        doc = Document(doc)
    idx = {}
    for p in doc.paragraphs:
        lv = _heading_level(p.style.name if p.style else None)
        if lv:
            idx.setdefault(p.text, []).append({"level": lv, "paragraph": p})
    return idx


def find_heading(doc, text, exact=True):
    """按标题文本找段落对象。exact=False 做包含匹配（取第一个命中）。"""
    if isinstance(doc, str):
        doc = Document(doc)
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            if exact and p.text == text:
                return p
            if not exact and text in p.text:
                return p
    return None


def find_paragraph_containing(doc, substring):
    """按正文文本子串查找段落（排除空段）。"""
    if isinstance(doc, str):
        doc = Document(doc)
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


# ================= 定点注入 =================

def insert_paragraph_after(target_p, new_text="", style_name=None):
    """在目标段后插一个段落，返回新 Paragraph。"""
    new_p = OxmlElement("w:p")
    target_p._p.addnext(new_p)
    wrap = Paragraph(new_p, target_p._parent)
    if style_name:
        wrap.style = style_name
    if new_text:
        wrap.add_run(new_text)
    return wrap


def insert_paragraph_before_cjk(target_p, new_text, name="宋体", size=12):
    """在目标段前插一个中文段落（自动 eastAsia 字体）。"""
    new_p = OxmlElement("w:p")
    target_p._p.addprevious(new_p)
    wrap = Paragraph(new_p, target_p._parent)
    run = wrap.add_run(new_text)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    return wrap


def insert_table_before(paragraph, rows=2, cols=2, caption=None, fill=None):
    """在目标段前插入真实表格（复制临时表 XML，无外链依赖）。
    fill: 可选填充函数 fill(ri, ci) -> str。返回 Table。"""
    tmp = Document()
    tbl = tmp.add_table(rows=rows, cols=cols)
    tbl.style = "Table Grid"
    for ri in range(rows):
        for ci in range(cols):
            tbl.rows[ri].cells[ci].text = fill(ri, ci) if fill else f"新{ri}{ci}"
    tbl_el = copy.deepcopy(tbl._tbl)
    if caption:
        p_new = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = caption
        r.append(t)
        p_new.append(r)
        paragraph._p.addprevious(p_new)
    paragraph._p.addprevious(tbl_el)
    return Table(tbl_el, paragraph._parent)


def insert_table_after(paragraph, rows=2, cols=2, caption=None, fill=None):
    """在目标段后插入表格（题注段在前、表格紧随其后的规范顺序）。"""
    tmp = Document()
    tbl = tmp.add_table(rows=rows, cols=cols)
    tbl.style = "Table Grid"
    for ri in range(rows):
        for ci in range(cols):
            tbl.rows[ri].cells[ci].text = fill(ri, ci) if fill else f"新{ri}{ci}"
    tbl_el = copy.deepcopy(tbl._tbl)
    anchor = paragraph._p
    if caption:
        p_new = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = caption
        r.append(t)
        p_new.append(r)
        anchor.addnext(p_new)
        anchor = p_new
    anchor.addnext(tbl_el)
    return Table(tbl_el, paragraph._parent)


# ================= 书签 & 超链接 =================

def list_bookmarks(doc):
    """全部书签：[{"name", "id"}]"""
    if isinstance(doc, str):
        doc = Document(doc)
    return [{"name": b.get(qn("w:name")), "id": b.get(qn("w:id"))}
            for b in doc.element.body.iter(qn("w:bookmarkStart"))]


def find_bookmark_paragraph(doc, name):
    """按名找书签所在段落 (Paragraph, bookmarkStart 元素)；找不到返回 (None, None)。"""
    if isinstance(doc, str):
        doc = Document(doc)
    for bm in doc.element.body.iter(qn("w:bookmarkStart")):
        if bm.get(qn("w:name")) == name:
            parent = bm.getparent()
            while parent is not None and parent.tag != qn("w:p"):
                parent = parent.getparent()
            return Paragraph(parent, doc), bm
    return None, None


def append_text_to_bookmark(doc, name, text):
    """在书签内容末尾追加文本（插到 bookmarkEnd 前，书签仍含新内容）。"""
    if isinstance(doc, str):
        doc = Document(doc)
    para, bm_start = find_bookmark_paragraph(doc, name)
    if bm_start is None:
        return False
    bm_id = bm_start.get(qn("w:id"))
    for el in para._p.iter():
        if el.tag == qn("w:bookmarkEnd") and el.get(qn("w:id")) == bm_id:
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = text
            r.append(t)
            el.addprevious(r)
            return True
    return False


def add_bookmark(paragraph, name, bookmark_id=0):
    """给段落加书签（包住整段现有文字），返回下一个可用 id。"""
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))
    p_el = paragraph._p
    p_el.insert(0, bm_start)
    p_el.append(bm_end)
    return bookmark_id + 1


def add_hyperlink(paragraph, url, text, color="0563C1", size_pt=11, underline=True):
    """段落末尾插入外部超链接（Word 默认蓝色下划线）。"""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    rPr.append(color_el)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)
    return hl


def list_hyperlinks(doc):
    """全部超链接：[(文本, target_url), ...]"""
    if isinstance(doc, str):
        doc = Document(doc)
    out = []
    for hl in doc.element.body.iter(qn("w:hyperlink")):
        rid = hl.get(qn("r:id"))
        target = doc.part.rels[rid].target_ref if rid in doc.part.rels else None
        htext = "".join(t.text or "" for t in hl.iter(qn("w:t")))
        out.append((htext, target))
    return out

__all__ = [
    "docx_to_json", "docx_to_markdown", "build_heading_index", "find_heading",
    "find_paragraph_containing", "insert_paragraph_after",
    "insert_paragraph_before_cjk", "insert_table_before", "insert_table_after",
    "list_bookmarks", "find_bookmark_paragraph", "append_text_to_bookmark",
    "add_bookmark", "add_hyperlink", "list_hyperlinks",
]
