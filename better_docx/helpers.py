# -*- coding: utf-8 -*-
"""
better_docx.helpers —— 基础助手函数：字体、段落、表格样式、图片插入、
全局替换，以及分页异常的诊断与修复。
"""
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

BODY_FONT = "宋体"
HEAD_FONT = "黑体"
CODE_FONT = "Consolas"


def set_font(run, name=BODY_FONT, size=11, bold=False, italic=False, color=None):
    """设置 run 字体。中文必须同时设 eastAsia，否则 Word 回退字体。
    color 可为 RGBColor 或 (r,g,b) 元组。"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        rgb = RGBColor(*color) if isinstance(color, tuple) else color
        run.font.color.rgb = rgb
    return run


def add_para(doc, text="", name=BODY_FONT, size=11, bold=False, italic=False,
             color=None, align=None, indent_cm=None, line_spacing=1.5, space_after=4):
    """添加正文段落，返回 Paragraph。text 为空则只建格式空段。"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if indent_cm:
        pf.first_line_indent = Cm(indent_cm)
    if line_spacing:
        pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    if text:
        set_font(p.add_run(text), name=name, size=size, bold=bold, italic=italic,
                 color=color)
    return p


def add_heading(doc, text, level=1):
    """标题：黑体加粗，一级 15pt / 二级 13pt / 三级 12pt。保留大纲级别（导航/目录可用）。"""
    h = doc.add_heading(text, level=level)
    size = {1: 15, 2: 13}.get(level, 12)
    for run in h.runs:
        set_font(run, name=HEAD_FONT, size=size, bold=True)
    return h


def shade_cell(cell, hex_fill):
    """单元格底纹。val 必须 clear（SOLID 会渲染全黑）。hex_fill 如 'D9E2F3'。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=11, name=BODY_FONT, color=None,
                  align=WD_ALIGN_PARAGRAPH.CENTER, vcenter=True):
    """清空单元格并写入文本，默认水平垂直居中。注意：会丢原格式。
    实现要点：cell.text="" 会留下「一个段落+一个空 run」，复用该空 run
    （runs[0].text=text），避免残留多余空 run。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    if vcenter:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    set_font(run, name=name, size=size, bold=bold, color=color)
    return p


def add_para_to_cell(cell, text, bold=False, size=11, name=BODY_FONT,
                     indent_cm=None, line_spacing=1.5, space_after=4):
    """在单元格末尾追加段落（保留已有内容），实验报告填表常用。"""
    p = cell.add_paragraph()
    pf = p.paragraph_format
    if indent_cm:
        pf.first_line_indent = Cm(indent_cm)
    if line_spacing:
        pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    set_font(p.add_run(text), name=name, size=size, bold=bold)
    return p


def new_image_paragraph(doc):
    """【关键】插图段落必须用 body 元素创建，否则 add_picture 报
    'CT_Body' object has no attribute 'part'。
    正确写法：doc.element.body.add_p()（CT_Body 方法）。
    坑：doc._body 是 _Body wrapper，没有 add_p() —— 会 AttributeError。"""
    return Paragraph(doc.element.body.add_p(), doc)


def add_image(doc, path, width_inches=5.5, caption=None, center=True):
    """正文插图 + 可选居中与题注。图片不存在时打印跳过，不抛异常。"""
    import os
    if not os.path.exists(path):
        print(f"[skip] 图片不存在: {path}")
        return None
    p = new_image_paragraph(doc)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_inches))
    if caption:
        add_para(doc, caption, size=10,
                 color=(0x60, 0x60, 0x60), italic=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0, space_after=6)
    return p


def add_image_to_cell(cell, path, width_inches=4.8, caption=None):
    """表格单元格内插图 + 可选居中题注（实验报告截图常用）。"""
    import os
    if not os.path.exists(path):
        print(f"[skip] 图片不存在: {path}")
        return None
    p_img = cell.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(path, width=Inches(width_inches))
    if caption:
        p_cap = cell.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p_cap.add_run(caption), size=10, italic=True,
                 color=(0x60, 0x60, 0x60))
    return p_img


def remove_images_from_paragraph(p):
    """移除段落里的所有图片（w:drawing）。替换模板旧图用。"""
    for drawing in p._element.findall(qn("w:drawing")):
        p._element.remove(drawing)


def global_replace(doc, old, new):
    """全局保格式替换：遍历正文段落 + 表格单元格内段落的所有 run。
    返回替换次数。限制：old 若被 Word 拆进多个 run 则匹配不到（先 merge_runs）。"""
    def do_paras(paras):
        n = 0
        for p in paras:
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    n += 1
        return n
    cnt = do_paras(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                cnt += do_paras(cell.paragraphs)
    return cnt


def set_cell_widths(table, widths_cm):
    """固定列宽：autofit=False + 逐格设宽。widths_cm 各列之和应 ≤ 正文宽(约16cm)。"""
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


def set_repeat_header(table, row_idx=0):
    """长表格跨页时每页重复表头行（w:tblHeader）。"""
    trPr = table.rows[row_idx]._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def set_row_no_split(row):
    """禁止该行跨页拆分（w:cantSplit）。"""
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    trPr.append(el)


# ================= 分页异常诊断与修复（子功能 10） =================

def diagnose_pagination(doc):
    """诊断"一页一段"分页异常。返回统计 dict：
    page_break_before_paragraphs / hard_page_breaks / hard_page_breaks_in_tables
    / styles_with_break_before（样式层面开启段前分页的段落样式名，信息性）。
    当结果非零时，文档很可能存在分页异常，可用 fix_one_paragraph_per_page 修复。"""
    pbb_count = 0
    hard_breaks = 0
    hard_in_tables = 0
    styles_with_pbb = []

    def scan(p, is_table=False):
        nonlocal pbb_count, hard_breaks, hard_in_tables
        pPr = p._p.pPr
        if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
            pbb_count += 1
        for r in p._p.findall(qn("w:r")):
            for br in r.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    hard_breaks += 1
                    if is_table:
                        hard_in_tables += 1

    for p in doc.paragraphs:
        scan(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    scan(p, is_table=True)

    for s in doc.styles:
        try:
            if s.type == WD_STYLE_TYPE.PARAGRAPH and s.paragraph_format.page_break_before:
                styles_with_pbb.append(s.name)
        except Exception:
            pass

    return {
        "page_break_before_paragraphs": pbb_count,
        "hard_page_breaks": hard_breaks,
        "hard_page_breaks_in_tables": hard_in_tables,
        "styles_with_break_before": styles_with_pbb,
    }


def fix_one_paragraph_per_page(doc, fix_page_break_before=True,
                               fix_hard_page_breaks=True, include_tables=True):
    """修复"一页一段"异常：关闭段落直接的段前分页 + 删除硬分页符 w:br type=page。
    实现要点（对比 2.txt 原版脚本的坑）：
    - 只删分页符元素（r.remove(br)），绝不 run.clear()——文字与分页符同 run 时
      clear() 会连文字一起误删，造成数据丢失。
    - doc.paragraphs 不含表格内段落，必须显式遍历 doc.tables 的 cells。
    - 样式层面的段前分页（如 Heading 1 常带）不动：那可能是排版意图。
    返回统计 dict。"""
    stats = {"page_break_before": 0, "hard_page_breaks": 0}

    def fix(p):
        if fix_page_break_before:
            pPr = p._p.pPr
            if pPr is not None:
                pbb = pPr.find(qn("w:pageBreakBefore"))
                if pbb is not None:
                    pPr.remove(pbb)
                    stats["page_break_before"] += 1
        if fix_hard_page_breaks:
            for r in p._p.findall(qn("w:r")):
                for br in r.findall(qn("w:br")):
                    if br.get(qn("w:type")) == "page":
                        r.remove(br)
                        stats["hard_page_breaks"] += 1

    for p in doc.paragraphs:
        fix(p)
    if include_tables:
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        fix(p)
    return stats

__all__ = [
    "set_font", "add_para", "add_heading", "shade_cell", "set_cell_text",
    "add_para_to_cell", "new_image_paragraph", "add_image", "add_image_to_cell",
    "remove_images_from_paragraph", "global_replace", "set_cell_widths",
    "set_repeat_header", "set_row_no_split", "diagnose_pagination",
    "fix_one_paragraph_per_page",
]
